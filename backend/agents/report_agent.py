"""
Report Agent — generates a structured multi-section markdown report via Gemini.
Optionally saves a .docx artifact using python-docx.
"""
import json
import os
import uuid
from pathlib import Path

import google.generativeai as genai

ARTIFACTS_DIR = Path(__file__).parent.parent / "artifacts"


def load_prompt(name: str) -> str:
    path = Path(__file__).parent.parent / "prompts" / f"{name}.txt"
    return path.read_text(encoding="utf-8")


async def run_report_agent(
    report_type: str,
    physician_list: str,
    sections: list[str] | None = None,
    icd10_context: str = "",
    geographic_scope: str = "",
) -> dict:
    try:
        physicians = json.loads(physician_list) if isinstance(physician_list, str) else physician_list

        genai.configure(api_key=os.environ["GOOGLE_API_KEY"])
        model = genai.GenerativeModel(
            model_name=os.environ.get("GEMINI_MODEL", "gemini-2.0-flash"),
            system_instruction=load_prompt("report_agent"),
        )

        sections_list = sections or [
            "Executive Summary",
            "Market Landscape",
            "Physician Segmentation",
            "Geographic Analysis",
            "Strategic Recommendations",
        ]

        prompt = (
            f"Report type: {report_type}\n"
            f"ICD-10 context: {icd10_context}\n"
            f"Geographic scope: {geographic_scope}\n"
            f"Sections to include: {', '.join(sections_list)}\n\n"
            f"Physician data (JSON):\n{json.dumps(physicians, indent=2)}\n\n"
            "Write the full report in markdown. Use ## for section headers. "
            "Reference physician names, specialties, states, and claim counts concretely."
        )

        import asyncio
        response = await asyncio.to_thread(model.generate_content, prompt)
        markdown = response.text.strip()

        # Optionally save .docx
        artifact_id = _save_docx(markdown, report_type)

        return {
            "markdown": markdown,
            "artifact_id": artifact_id,
            "filename": "report.docx",
        }

    except Exception as exc:
        return {"error": True, "message": str(exc)}


def _save_docx(markdown: str, title: str) -> str:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=0)

    for line in markdown.split("\n"):
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)

    artifact_id = str(uuid.uuid4())
    out_path = ARTIFACTS_DIR / f"{artifact_id}.docx"
    doc.save(str(out_path))
    return artifact_id
